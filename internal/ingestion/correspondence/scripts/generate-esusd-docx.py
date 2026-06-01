"""
SkaFld Trailhead — ESUSD Implementation Plan (DOCX)

Uses the branded Foundry document template to preserve header (logo top-right,
orange left stripe), styles, and numbering definitions, then replaces all body
content with the implementation plan.

Typography (from SkaFld Brand Guidelines):
  - Display / Headlines:  Manrope (geometric sans-serif, extra bold)
  - Body / UI:            Inter (sans-serif)

Color Palette (from skafldstudio.com):
  - Primary Orange:   #E07800
  - Primary Dark:     #1A1A1A
  - Body Text:        #333333
  - Gray:             #919191
  - Cream:            #FAF6F1
  - Accent Line:      #F0E6D9
  - Header BG:        #2D2D2D
"""

import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

# ── Brand Typography ────────────────────────────────────────────────
DISPLAY_FONT = 'Manrope'
HEADING_FONT = 'Manrope'
BODY_FONT    = 'Inter'

# ── Brand Colors ────────────────────────────────────────────────────
SKAFLD_ORANGE = RGBColor(0xE0, 0x78, 0x00)
BRAND_DARK    = RGBColor(0x1A, 0x1A, 0x1A)
BODY_COLOR    = RGBColor(0x33, 0x33, 0x33)
SUBTITLE_GRAY = RGBColor(0x55, 0x55, 0x55)
MUTED_GRAY    = RGBColor(0x88, 0x88, 0x88)
GRAY          = RGBColor(0x91, 0x91, 0x91)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)

HEX_ORANGE     = 'E07800'
HEX_DARK       = '1A1A1A'
HEX_CREAM      = 'FAF6F1'
HEX_WHITE      = 'FFFFFF'
HEX_LIGHT_GRAY = 'F5F5F5'
HEX_HEADER_BG  = '2D2D2D'
HEX_ACCENT     = 'F0E6D9'

# ── Type Sizes ──────────────────────────────────────────────────────
BODY_SIZE         = Pt(11)
H1_SIZE           = Pt(20)
H2_SIZE           = Pt(16)
H3_SIZE           = Pt(13)
TITLE_SIZE        = Pt(24)
TABLE_SIZE        = Pt(9)
TABLE_HEADER_SIZE = Pt(9)

# ── Spacing ─────────────────────────────────────────────────────────
SP_NONE    = Pt(0)
SP_TIGHT   = Pt(2)
SP_NORMAL  = Pt(3)
SP_SECTION = Pt(6)
SP_HEADING = Pt(10)
SP_MAJOR   = Pt(14)

SCRIPT_DIR    = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(
    os.path.expanduser('~'),
    'Development/projects-skafld/skafld-foundry/foundry-internal/assets/templates',
    'SkaFld_Foundry_Document_Template.docx'
)
OUTPUT_DIR    = os.path.join(SCRIPT_DIR, '..', 'esusd-partnership')
OUTPUT_PATH   = os.path.join(OUTPUT_DIR, 'ESUSD_Implementation_Plan.docx')

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

BULLET_NUM_ID  = 1


# ═════════════════════════════════════════════════════════════════════
# LOW-LEVEL HELPERS
# ═════════════════════════════════════════════════════════════════════

def _add_numpr(paragraph, num_id=1, ilvl=0):
    pPr = paragraph._element.get_or_add_pPr()
    numPr = etree.SubElement(pPr, f'{{{WNS}}}numPr')
    ilvl_el = etree.SubElement(numPr, f'{{{WNS}}}ilvl')
    ilvl_el.set(f'{{{WNS}}}val', str(ilvl))
    numId_el = etree.SubElement(numPr, f'{{{WNS}}}numId')
    numId_el.set(f'{{{WNS}}}val', str(num_id))


def _spacing(paragraph, before=None, after=None, line_spacing=None):
    pf = paragraph.paragraph_format
    if before is not None:
        pf.space_before = before
    if after is not None:
        pf.space_after = after
    if line_spacing is not None:
        pf.line_spacing = line_spacing


def _run(paragraph, text, size=BODY_SIZE, bold=None, italic=None, font=None, color=None):
    r = paragraph.add_run(text)
    r.font.name = font or BODY_FONT
    r.font.size = size
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return r


def _update_heading_styles(doc):
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = HEADING_FONT
    h2_style.font.size = H1_SIZE
    h2_style.font.bold = True
    h2_style.font.italic = False
    h2_style.font.color.rgb = BRAND_DARK
    h2_style.paragraph_format.space_before = SP_MAJOR
    h2_style.paragraph_format.space_after = SP_NORMAL

    h3_style = doc.styles['Heading 3']
    h3_style.font.name = BODY_FONT
    h3_style.font.size = H2_SIZE
    h3_style.font.bold = True
    h3_style.font.italic = False
    h3_style.font.color.rgb = BRAND_DARK
    h3_style.paragraph_format.space_before = SP_HEADING
    h3_style.paragraph_format.space_after = SP_NORMAL

    normal_style = doc.styles['Normal']
    normal_style.font.name = BODY_FONT
    normal_style.font.size = BODY_SIZE
    normal_style.font.bold = False
    normal_style.font.italic = False
    normal_style.font.color.rgb = BODY_COLOR

    if 'p1' in [s.name for s in doc.styles]:
        p1_style = doc.styles['p1']
        p1_style.font.name = BODY_FONT
        p1_style.font.size = BODY_SIZE
        p1_style.font.color.rgb = BODY_COLOR

    if 'List Paragraph' in [s.name for s in doc.styles]:
        lp_style = doc.styles['List Paragraph']
        lp_style.font.name = BODY_FONT
        lp_style.font.size = BODY_SIZE
        lp_style.font.color.rgb = BODY_COLOR


def _clear_body(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith('}p') or child.tag.endswith('}tbl'):
            body.remove(child)


def _style_name(doc, preferred='p1'):
    return preferred if preferred in [s.name for s in doc.styles] else 'Normal'


# ── Table Helpers ───────────────────────────────────────────────────

def _set_cell_shading(cell, color_hex):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shading = etree.SubElement(tcPr, qn('w:shd'))
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)


def _set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = etree.SubElement(tcPr, qn('w:tcBorders'))
    for edge, spec in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = etree.SubElement(tcBorders, qn(f'w:{edge}'))
        if spec:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(spec[1]))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), spec[0])
        else:
            el.set(qn('w:val'), 'none')
            el.set(qn('w:sz'), '0')
            el.set(qn('w:space'), '0')


def _set_cell_padding(cell, top=40, bottom=40, left=80, right=80):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    mar = etree.SubElement(tcPr, qn('w:tcMar'))
    for edge, val in [('top', top), ('bottom', bottom), ('start', left), ('end', right)]:
        el = etree.SubElement(mar, qn(f'w:{edge}'))
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')


def _set_cell_width(cell, width_inches):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcW = etree.SubElement(tcPr, qn('w:tcW'))
    tcW.set(qn('w:w'), str(int(width_inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')


def _set_cell_vertical_alignment(cell, align='center'):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:vAlign')):
        tcPr.remove(old)
    vAlign = etree.SubElement(tcPr, qn('w:vAlign'))
    vAlign.set(qn('w:val'), align)


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else etree.SubElement(tbl, qn('w:tblPr'))
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblBorders = etree.SubElement(tblPr, qn('w:tblBorders'))
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = etree.SubElement(tblBorders, qn(f'w:{edge}'))
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')


def _set_table_width(table, width_pct=100):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else etree.SubElement(tbl, qn('w:tblPr'))
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW = etree.SubElement(tblPr, qn('w:tblW'))
    tblW.set(qn('w:w'), str(width_pct * 50))
    tblW.set(qn('w:type'), 'pct')


def _format_cell_text(cell, text, size=TABLE_SIZE, bold=False, italic=False,
                      font=None, color=None, align=WD_PARAGRAPH_ALIGNMENT.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    _spacing(p, before=SP_NONE, after=SP_NONE, line_spacing=1.2)
    r = p.add_run(text)
    r.font.name = font or BODY_FONT
    r.font.size = size
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color or BODY_COLOR
    return r


def _prevent_table_split(table):
    """Prevent any row from splitting across pages, and keep all rows together.

    Sets w:cantSplit on every row and w:keepNext on every row except the last,
    so the entire table stays on a single page where possible.
    """
    rows = list(table.rows)
    for i, row in enumerate(rows):
        tr = row._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = etree.SubElement(tr, qn('w:trPr'))
            tr.insert(0, trPr)
        for tag in ('w:cantSplit', 'w:keepNext'):
            for old in trPr.findall(qn(tag)):
                trPr.remove(old)
        cantSplit = etree.SubElement(trPr, qn('w:cantSplit'))
        cantSplit.set(qn('w:val'), 'true')
        if i < len(rows) - 1:
            keepNext = etree.SubElement(trPr, qn('w:keepNext'))
            keepNext.set(qn('w:val'), 'true')


def add_branded_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    n_rows = len(rows)
    col_alignments = [WD_PARAGRAPH_ALIGNMENT.LEFT] * n_cols
    if col_widths is None:
        page_width = 6.5
        col_widths = [page_width / n_cols] * n_cols

    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    _remove_table_borders(table)
    _set_table_width(table)

    # Header row
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        _set_cell_shading(cell, HEX_HEADER_BG)
        _set_cell_padding(cell, top=50, bottom=50, left=70, right=70)
        _set_cell_vertical_alignment(cell, 'center')
        _set_cell_width(cell, col_widths[ci])
        _set_cell_borders(cell,
                          top=(HEX_ORANGE, 6),
                          bottom=(HEX_ORANGE, 6),
                          left=None, right=None)
        _format_cell_text(cell, h.upper(),
                          size=TABLE_HEADER_SIZE, bold=True,
                          font=HEADING_FONT, color=WHITE,
                          align=col_alignments[ci])

    # Body rows
    for ri, row_data in enumerate(rows):
        is_last = ri == n_rows - 1
        bg = HEX_WHITE if ri % 2 == 0 else HEX_LIGHT_GRAY

        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            _set_cell_shading(cell, bg)
            _set_cell_padding(cell, top=35, bottom=35, left=70, right=70)
            _set_cell_vertical_alignment(cell, 'center')
            _set_cell_width(cell, col_widths[ci])

            bottom_border = (HEX_DARK, 8) if is_last else (HEX_ACCENT, 2)
            _set_cell_borders(cell, top=None, bottom=bottom_border, left=None, right=None)

            _format_cell_text(cell, str(val), size=TABLE_SIZE, color=BODY_COLOR,
                              align=col_alignments[ci])

    _prevent_table_split(table)
    return table


# ── Paragraph Helpers ───────────────────────────────────────────────

def add_title(doc, text):
    try:
        p = doc.add_paragraph(text, style='Title')
    except KeyError:
        p = doc.add_paragraph()
        _run(p, text, TITLE_SIZE, bold=True, font=DISPLAY_FONT, color=BRAND_DARK)
    return p


def add_h1(doc, text, page_break_before=False):
    p = doc.add_paragraph(text, style='Heading 2')
    p.paragraph_format.keep_with_next = True
    if page_break_before:
        p.paragraph_format.page_break_before = True
    return p


def add_h2(doc, text):
    p = doc.add_paragraph(text, style='Heading 3')
    p.paragraph_format.keep_with_next = True
    return p


def add_h3(doc, text):
    p = doc.add_paragraph(style=_style_name(doc))
    _run(p, text, H3_SIZE, bold=True, font=BODY_FONT, color=SKAFLD_ORANGE)
    _spacing(p, before=SP_SECTION, after=SP_TIGHT)
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text, space_before=None, space_after=None):
    p = doc.add_paragraph(style=_style_name(doc))
    _run(p, text)
    _spacing(p,
             before=space_before if space_before is not None else SP_NORMAL,
             after=space_after if space_after is not None else SP_NORMAL,
             line_spacing=1.15)
    return p


def add_bold_body(doc, label, text):
    p = doc.add_paragraph(style=_style_name(doc))
    _run(p, label, bold=True)
    if text:
        _run(p, text)
    _spacing(p, before=SP_NORMAL, after=SP_NORMAL, line_spacing=1.15)
    return p


def add_bullet(doc, label, text='', num_id=BULLET_NUM_ID):
    p = doc.add_paragraph(style='List Paragraph')
    _add_numpr(p, num_id=num_id, ilvl=0)
    if text:
        _run(p, label, bold=True)
        _run(p, text)
    else:
        _run(p, label)
    _spacing(p, before=SP_TIGHT, after=SP_TIGHT, line_spacing=1.15)
    return p


# ═════════════════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ═════════════════════════════════════════════════════════════════════

L = WD_PARAGRAPH_ALIGNMENT.LEFT


def _add_page_number(doc):
    """Add page numbers to the footer of every section."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        _spacing(p, before=SP_NONE, after=SP_NONE)
        r = p.add_run()
        r.font.name = BODY_FONT
        r.font.size = Pt(9)
        r.font.color.rgb = GRAY
        fldChar1 = etree.SubElement(r._element, qn('w:fldChar'))
        fldChar1.set(qn('w:fldCharType'), 'begin')
        r2 = p.add_run()
        r2.font.name = BODY_FONT
        r2.font.size = Pt(9)
        r2.font.color.rgb = GRAY
        instrText = etree.SubElement(r2._element, qn('w:instrText'))
        instrText.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        instrText.text = ' PAGE '
        r3 = p.add_run()
        r3.font.name = BODY_FONT
        r3.font.size = Pt(9)
        r3.font.color.rgb = GRAY
        fldChar2 = etree.SubElement(r3._element, qn('w:fldChar'))
        fldChar2.set(qn('w:fldCharType'), 'end')


def _set_margins(doc, top=0.9, bottom=0.75, left=1.0, right=1.0):
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def _resize_header_logo(doc, target_width_emu=939600, target_height_emu=180000):
    """Resize the header logo to ~1in x 0.2in."""
    WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    for section in doc.sections:
        hdr = section.header
        if hdr is None:
            continue
        hdr_xml = hdr._element
        for ext in hdr_xml.iter(f'{{{WP_NS}}}extent'):
            cx = int(ext.get('cx', 0))
            cy = int(ext.get('cy', 0))
            if cy > 5000000:
                continue
            ext.set('cx', str(target_width_emu))
            ext.set('cy', str(target_height_emu))
        for ext in hdr_xml.iter(f'{{{A_NS}}}ext'):
            cx = int(ext.get('cx', 0))
            cy = int(ext.get('cy', 0))
            if cy > 5000000:
                continue
            if cx > 100000:
                ext.set('cx', str(target_width_emu))
                ext.set('cy', str(target_height_emu))


def _narrow_header_stripe(doc):
    """Make the orange left stripe narrower."""
    WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    for section in doc.sections:
        hdr = section.header
        if hdr is None:
            continue
        hdr_xml = hdr._element
        for ext in hdr_xml.iter(f'{{{WP_NS}}}extent'):
            cx = int(ext.get('cx', 0))
            cy = int(ext.get('cy', 0))
            # The stripe is the very tall, narrow element
            if cy > 5000000:
                # Narrow it: ~3pt width (38100 EMU = ~0.04in)
                ext.set('cx', str(38100))
        for ext in hdr_xml.iter(f'{{{A_NS}}}ext'):
            cx = int(ext.get('cx', 0))
            cy = int(ext.get('cy', 0))
            if cy > 5000000:
                ext.set('cx', str(38100))


def _reposition_header_stripe(doc, h_offset_emu=-992525):
    """Move the orange stripe to the left page edge."""
    WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    for section in doc.sections:
        hdr = section.header
        if hdr is None:
            continue
        for anchor in hdr._element.iter(f'{{{WP_NS}}}anchor'):
            posH = anchor.find(f'{{{WP_NS}}}positionH')
            if posH is not None:
                offset = posH.find(f'{{{WP_NS}}}posOffset')
                if offset is not None:
                    offset.text = str(h_offset_emu)
                    break


def add_detail_line(doc, label, value):
    """Add a label: value line for the title block metadata."""
    p = doc.add_paragraph(style=_style_name(doc))
    _run(p, label, size=Pt(9), bold=True, font=BODY_FONT, color=GRAY)
    _run(p, '  ' + value, size=Pt(9), font=BODY_FONT, color=BODY_COLOR)
    _spacing(p, before=Pt(1), after=Pt(1), line_spacing=1.3)
    return p


def add_orange_rule(doc):
    """Add a thin orange horizontal rule."""
    p = doc.add_paragraph(style=_style_name(doc))
    _spacing(p, before=SP_SECTION, after=SP_SECTION)
    pPr = p._element.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn('w:pBdr'))
    bottom = etree.SubElement(pBdr, qn('w:bottom'))
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), HEX_ORANGE)
    return p


def create_document():
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f'Template not found: {TEMPLATE_PATH}')

    doc = docx.Document(TEMPLATE_PATH)
    _update_heading_styles(doc)

    # Professional margins & page numbers
    _set_margins(doc, top=0.9, bottom=0.75, left=1.0, right=1.0)

    # Update footer styling
    for section in doc.sections:
        footer = section.footer
        if footer is not None:
            for paragraph in footer.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = BODY_FONT
                    run.font.italic = True
                    run.font.color.rgb = SKAFLD_ORANGE

    _add_page_number(doc)
    _resize_header_logo(doc)
    _narrow_header_stripe(doc)
    _reposition_header_stripe(doc)

    _clear_body(doc)

    # ── TITLE BLOCK ──────────────────────────────────────────────
    add_title(doc, 'SkaFld Trailhead')

    # Subtitle
    p = doc.add_paragraph(style=_style_name(doc))
    _run(p, 'ESUSD Partnership Implementation Plan',
         size=Pt(14), font=DISPLAY_FONT, color=SUBTITLE_GRAY)
    _spacing(p, before=SP_NONE, after=SP_HEADING)

    # Thin orange divider
    add_orange_rule(doc)

    # Metadata lines (no table, clean layout)
    add_detail_line(doc, 'Prepared By:', 'SkaFld Studio LLC')
    add_detail_line(doc, 'Prepared For:', 'El Segundo Unified School District')
    add_detail_line(doc, 'Contact:', 'Charles Sims — charles@skafldstudio.com')
    add_detail_line(doc, 'Date:', 'March 2026')

    # ── OVERVIEW ─────────────────────────────────────────────────
    add_h1(doc, 'Overview')
    add_body(doc,
        'SkaFld Trailhead is a portfolio-based career pathway program where high school '
        'students build employer-validated portfolios by solving real problems for local '
        'companies using AI. Students develop durable metacognitive skills — decomposition '
        'thinking, cross-domain pattern recognition, and responsible AI practices — that '
        'transfer across any tool or domain.')
    add_body(doc,
        'This plan outlines the implementation partnership between SkaFld Studio LLC and '
        'El Segundo Unified School District to pilot SkaFld Trailhead with 12 students '
        'over 12 months.')

    # ── PROGRAM AT A GLANCE ──────────────────────────────────────
    add_h2(doc, 'Program at a Glance')
    add_branded_table(doc,
        headers=['Element', 'Detail'],
        rows=[
            ('Students', '12 in 1 cross-grade team (9th-12th)'),
            ('Teacher Mentor', '1 dedicated teacher mentor'),
            ('Session Format', 'Monthly 90-minute after-school sessions'),
            ('Employer Partners', '2-3 local companies (aerospace, tech)'),
            ('Curriculum', 'SkaFld Ideation Methodology'),
            ('Platform', 'SkaFld Trailhead (AI-powered, COPPA/FERPA compliant)'),
            ('Duration', '12 months'),
        ],
        col_widths=[2.0, 4.5],
    )

    # ── RESPONSIBILITY MATRIX ────────────────────────────────────
    add_h1(doc, 'Responsibility Matrix', page_break_before=True)

    add_h3(doc, 'What SkaFld Provides')
    add_bullet(doc, 'SkaFld Trailhead Platform', ' — AI Coaching Engine, Portfolio Review System, Student Progression Tracking, Teacher Dashboard. Fully COPPA/FERPA compliant with enterprise SSO (Microsoft/Google).')
    add_bullet(doc, 'Curriculum & Training', ' — Complete SkaFld Ideation Methodology curriculum for monthly 90-minute after-school sessions. Teacher mentor training and ongoing support.')
    add_bullet(doc, 'Employer Partnerships', ' — Recruitment and onboarding of 2-3 employer partners in El Segundo\'s aerospace/tech corridor for project briefs and portfolio review.')
    add_bullet(doc, 'Program Management', ' — Day-to-day coordination, progress tracking, employer communications, quarterly portfolio review facilitation.')
    add_bullet(doc, 'Assessment & Reporting', ' — Learning outcome tracking with disaggregated equity monitoring. Progress reports and year-end assessment.')

    add_h3(doc, 'What ESUSD Provides')
    add_bullet(doc, 'District Point of Contact', ' — A designated coordinator (CTE coordinator, assistant principal, or equivalent) to serve as the primary liaison between SkaFld and the district.')
    add_bullet(doc, 'Facility Access', ' — A classroom or meeting space for monthly 90-minute after-school sessions.')
    add_bullet(doc, 'Teacher Mentor', ' — 1 teacher mentor to facilitate team sessions. SkaFld provides all training and curriculum materials; the teacher brings their knowledge of students and school culture.')
    add_bullet(doc, 'Student Recruitment Support', ' — Help promoting the program to students and families and processing parental consent forms.')
    add_bullet(doc, 'Administrative Support', ' — Standard district approvals, data sharing agreements, and coordination with existing CTE or after-school programming.')

    # ── IMPLEMENTATION TIMELINE ──────────────────────────────────
    add_h1(doc, 'Implementation Timeline')

    add_h3(doc, 'Phase 1: Foundation (Months 1-3)')
    add_branded_table(doc,
        headers=['Month', 'SkaFld Actions', 'ESUSD Actions'],
        rows=[
            ('1', 'Complete SkaFld Trailhead platform buildout; begin employer partner recruitment',
                  'Designate district point of contact; identify facility for sessions'),
            ('2', 'Finalize platform (AI Coaching Engine, Portfolio Review System, Teacher Dashboard); onboard first employer partners',
                  'Identify teacher mentor; begin internal communications about program'),
            ('3', 'Platform ready for pilot; 2-3 employer partners signed; develop teacher training materials',
                  'Teacher mentor confirmed; facility reserved; student recruitment materials distributed'),
        ],
        col_widths=[0.6, 3.0, 2.9],
    )
    add_bold_body(doc, 'Month 3 Milestones:', '')
    add_bullet(doc, 'SkaFld Trailhead platform ready for pilot')
    add_bullet(doc, '2-3 signed employer partner commitments')
    add_bullet(doc, 'Teacher mentor identified and scheduled for training')

    add_h3(doc, 'Phase 2: Launch (Months 4-6)')
    add_branded_table(doc,
        headers=['Month', 'SkaFld Actions', 'ESUSD Actions'],
        rows=[
            ('4', 'Conduct teacher mentor training; finalize student applications; configure platform for ESUSD SSO',
                  'Open student applications; support equitable recruitment; process parental consent forms'),
            ('5', 'Form team of 12 students; assign employer project brief; launch monthly after-school sessions',
                  'Students begin attending sessions; teacher mentor actively facilitating'),
            ('6', 'First project cycle underway; mid-year progress report',
                  'Ongoing facility and schedule support; feedback on program integration'),
        ],
        col_widths=[0.6, 3.0, 2.9],
    )
    add_bold_body(doc, 'Month 5 Milestone: ', 'Pilot launched — 12 students in 1 team')
    add_bold_body(doc, 'Month 6 Milestone: ', 'Mid-year progress report completed')

    add_h3(doc, 'Phase 3: Execution (Months 7-10)')
    add_branded_table(doc,
        headers=['Month', 'SkaFld Actions', 'ESUSD Actions'],
        rows=[
            ('7-8', 'First project deliverables completed and submitted to employers',
                    'Teacher mentor supports student presentations'),
            ('9', 'First employer portfolio review panel — all 12 students receive employer feedback',
                  'Facilitate student participation in review panel'),
            ('10', 'Second project cycle launched with new employer briefs',
                   'Continue session support; share early outcomes with district leadership'),
        ],
        col_widths=[0.6, 3.0, 2.9],
    )
    add_bold_body(doc, 'Month 9 Milestone: ', 'First employer portfolio review cycle completed')

    add_h3(doc, 'Phase 4: Validation (Months 11-12)')
    add_branded_table(doc,
        headers=['Month', 'SkaFld Actions', 'ESUSD Actions'],
        rows=[
            ('11', 'Facilitate micro-internship placements (3+ students); second portfolio review',
                   'Support work-based learning logistics (permissions, scheduling)'),
            ('12', 'Year-end assessment; outcomes report; Year 2 expansion planning',
                   'Joint review of outcomes; discuss Year 2 continuation and expansion'),
        ],
        col_widths=[0.6, 3.0, 2.9],
    )
    add_bold_body(doc, 'Month 11 Milestone: ', '3+ micro-internship placements')
    add_bold_body(doc, 'Month 12 Milestone: ', 'Year-end report with outcomes and Year 2 plan')

    # ── IMMEDIATE NEXT STEPS ─────────────────────────────────────
    add_h1(doc, 'Immediate Next Steps (April 2026)')
    add_branded_table(doc,
        headers=['#', 'Action', 'Owner', 'Target Date'],
        rows=[
            ('1', 'Schedule 30-minute partnership kickoff call', 'Charles Sims / Dr. Johnson', 'Week of April 13'),
            ('2', 'Designate ESUSD point of contact', 'Dr. Johnson', 'April 18'),
            ('3', 'Identify facility for monthly after-school sessions', 'ESUSD Point of Contact', 'April 25'),
            ('4', 'Share teacher mentor criteria and identify candidate', 'SkaFld + ESUSD', 'May 2'),
            ('5', 'Execute data sharing / partnership agreement', 'Both parties', 'May 9'),
        ],
        col_widths=[0.4, 2.5, 1.8, 1.8],
    )

    # ── TEAM MODEL ───────────────────────────────────────────────
    add_h1(doc, 'Team Model')
    add_body(doc,
        'The cross-grade team model (12th graders mentoring 10th graders, 11th graders '
        'mentoring 9th graders) ensures peer support structures are embedded from day one. '
        'All platform features are built with accessibility in mind.')

    # ── KEY CONTACTS ─────────────────────────────────────────────
    add_h1(doc, 'Key Contacts', page_break_before=True)

    add_h3(doc, 'SkaFld Studio LLC')
    add_branded_table(doc,
        headers=['Name', 'Role', 'Email'],
        rows=[
            ('Charles Sims', 'Project Director', 'charles@skafldstudio.com'),
            ('Keith Coleman', 'Strategy & Partnerships', 'keith@skafldstudio.com'),
            ('Mike Belloli', 'Technology & Operations', 'mike@skafldstudio.com'),
            ('Danielle Joseph', 'Chief of Staff', 'danielle@skafldstudio.com'),
        ],
        col_widths=[2.0, 2.2, 2.3],
    )

    add_h3(doc, 'El Segundo Unified School District')
    add_branded_table(doc,
        headers=['Name', 'Role', 'Email'],
        rows=[
            ('Dr. Jason Johnson', 'Superintendent', 'jjohnson@esusd.net'),
            ('TBD', 'District Point of Contact', 'TBD'),
        ],
        col_widths=[2.0, 2.2, 2.3],
    )

    add_body(doc, '', space_before=SP_SECTION)
    p = add_body(doc, '', space_before=SP_NONE)
    _run(p, 'SkaFld Studio LLC — 889 N Douglas St, Suite 201, El Segundo, CA 90245',
         size=Pt(9), italic=True, color=GRAY)

    return doc


# ═════════════════════════════════════════════════════════════════════
# MAIN
# ═════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    print('Generating ESUSD Implementation Plan DOCX...')
    print(f'  Template: {TEMPLATE_PATH}')
    doc = create_document()
    doc.save(OUTPUT_PATH)
    print(f'  -> {OUTPUT_PATH}')
    print('Done!')
